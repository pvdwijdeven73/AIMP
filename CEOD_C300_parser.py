from html import unescape
import pandas as pd
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from os import system
from routines import dprint
from glob import glob


def get_params(line):
    line = unescape(line)
    curpos = 0
    params = {}
    target = "<ParamValue>"
    while curpos != -1:
        if target == "<ParamValue>":
            target = "<ParamName>"
        else:
            target = "<ParamValue>"
        target_end = "</" + target[1:]
        curpos = line.find(target, curpos)
        if curpos != -1:
            curpos += len(target)
            if target == "<ParamValue>":
                value = line[curpos : curpos + line[curpos:].find(target_end)]
                params[name] = value
            else:
                name = line[curpos : curpos + line[curpos:].find(target_end)]
    return params


def get_sequence(path, filename):
    f = open(path + filename, "r", encoding="utf-16")
    lines = f.readlines()
    target = ""
    blockname = ""

    TR_dict = {}
    STEP_dict = {}
    # for line in lines:
    line = "".join(lines)
    curpos = 0
    while curpos != -1:
        if target == "" or target == "<Parameters>" or blockname == "":
            target = "<BlockName>"
        elif target == "<ClassName>":
            target = "<Parameters>"
        else:
            target = "<ClassName>"
        target_end = "</" + target[1:]
        # print(f"{target} - {target_end}")
        curpos = line.find(target, curpos)
        if curpos != -1:
            curpos += len(target)
            if target == "<BlockName>":
                blockname = line[curpos : curpos + line[curpos:].find(target_end)]
                # print(f"BlockName: {blockname}")
            elif target == "<ClassName>":
                classname = line[curpos : curpos + line[curpos:].find(target_end)]
                # print(f"ClassName: {classname}")
            else:
                params = line[curpos : curpos + line[curpos:].find(target_end)]
                pars = get_params(params)
                if classname == "TRANSITION":
                    TR_dict[blockname] = pars
                elif classname == "STEP":
                    STEP_dict[blockname] = pars

    df_TR = pd.DataFrame(TR_dict)
    df_STEP = pd.DataFrame(STEP_dict)
    return df_TR, df_STEP, TR_dict, STEP_dict


def write_sequence_excel(df_TR, df_STEP, path, filename_output):
    with pd.ExcelWriter(path + filename_output, mode="w") as writer:
        df_TR.to_excel(writer, sheet_name="TR", index=True)
        df_STEP.to_excel(writer, sheet_name="STEP", index=True)
    return


def write_sequence_word(TR_dict, STEP_dict, path, filename_output):
    doc = docx.Document()
    doc.add_heading(filename_output.replace(".docx", ""), 0)
    style = doc.styles["Normal"]
    style.paragraph_format.line_spacing = 1
    style.paragraph_format.space_before = 0
    style.paragraph_format.space_after = 0
    font = style.font
    font.name = "Calibri"
    font.size = Pt(8)
    # font_styles = doc.styles
    # font_charstyle = font_styles.add_style("CommentsStyle", WD_STYLE_TYPE.CHARACTER)
    # font_object = font_charstyle.font
    # font_object.size = Pt(6)
    # font_object.name = "Calibri"

    # STEPS
    for step in sorted(STEP_dict.keys()):
        step_lst = []
        # print(f"{step} - {STEP_dict[step]['DESC']}")
        try:
            step_desc = STEP_dict[step]["DESC"].replace('"', "")
        except:
            step_desc = "NO STEP DESC FOUND!!!!"
        OP = 1
        while f"OP[{OP}].INSTRUCTTYPE" in STEP_dict[step]:
            instr = STEP_dict[step][f"OP[{OP}].INSTRUCTTYPE"].replace('"', "")
            if instr == "Info":
                # operator message found
                expr = STEP_dict[step][f"OP[{OP}].INSTRUCTION"].replace('"', "")
                desc = STEP_dict[step][f"OP[{OP}].DESC"].replace('"', "")
                delay = int(STEP_dict[step][f"OP[{OP}].DELAYTIME"])
            elif instr == "None":
                # expression found
                expr = STEP_dict[step][f"OP[{OP}].SRCEXPR"].replace('"', "")
                desc = STEP_dict[step][f"OP[{OP}].DESC"].replace('"', "")
                delay = int(STEP_dict[step][f"OP[{OP}].DELAYTIME"])
            else:
                dprint(f"Strange instruction type found {instr}", "RED")
            if delay > 0:
                delay = f" - {delay} seconds delay"
            else:
                delay = ""
            # print(f"{expr} - {desc} - #{OP}{delay}")
            step_lst.append([expr, desc, f"#{OP}{delay}"])
            OP += 1
        doc_para = doc.add_paragraph("\n")
        runner = doc_para.add_run(f"{step.split('.')[1]}")
        runner.bold = True
        doc_para.add_run(f' - "{step_desc}"')
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        row = table.rows[0].cells
        row[0].text = "Expression"
        row[1].text = "Description"
        row[2].text = "OP#"
        for expr, desc, id in step_lst:
            # Adding a row and then adding data in it.
            row = table.add_row().cells
            # Converting id to string as table can only take string input
            row[0].text = expr
            row[1].text = desc
            row[2].text = id
        table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[2].paragraphs[0].runs[0].font.bold = True
        for row in table.rows:
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            row.cells[0].paragraphs[0].runs[0].font.name = "Consolas"
        table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
        table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[2].paragraphs[0].runs[0].font.bold = True
        for cell in table.columns[0].cells:
            cell.width = Cm(7.25)
        for cell in table.columns[1].cells:
            cell.width = Cm(5.75)
        for cell in table.columns[2].cells:
            cell.width = Cm(2.8)

    # TRANSITIONS
    for tr in sorted(TR_dict.keys()):
        tr_lst = []
        try:
            tr_desc = TR_dict[tr]["DESC"].replace('"', "")
        except:
            tr_desc = "NO TRANSITION DESC FOUND!!!!"
        # print(f"{tr} - {TR_dict[tr]['DESC']}")
        OP = 1
        G1 = TR_dict[tr][f"G[1].ALGID"].replace('"', "")
        if G1 == "Connect":
            G1 = ""
        G2 = TR_dict[tr][f"G[2].ALGID"].replace('"', "")
        if G2 == "Connect":
            G2 = ""
        G3 = TR_dict[tr][f"G[3].ALGID"].replace('"', "")
        if G3 == "Connect":
            G3 = ""
        G4 = TR_dict[tr][f"G[4].ALGID"].replace('"', "")
        if G4 == "Connect":
            G4 = ""
        while f"C[{OP}].EXPR" in TR_dict[tr]:
            expr = TR_dict[tr][f"C[{OP}].EXPR"].replace('"', "")
            if expr != "":
                # expression found
                desc = TR_dict[tr][f"C[{OP}].DESC"].replace('"', "")
                gate = TR_dict[tr][f"C[{OP}].GATEASGN"].replace('"', "")
                if gate == "GateP1":
                    gate1 = G2
                elif gate == "GateP2":
                    gate1 = G3
                elif gate == "GateP3":
                    gate1 = G4
                else:
                    gate1 = ""
                gate2 = G1
            # print(f"{expr} - {desc} - {gate1} - {gate2}")
            try:
                tr_lst.append([expr, desc, gate1, gate2])
            except:
                dprint(
                    f'Transition {tr} of sequence {filename_output.replace(".docx", "")} is empty',
                    "RED",
                )
            OP += 1

        doc_para = doc.add_paragraph("\n")
        runner = doc_para.add_run(f"{tr.split('.')[1]}")
        runner.bold = True
        doc_para.add_run(f' - "{tr_desc}"')
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        row = table.rows[0].cells
        row[0].text = "Expression"
        row[1].text = "Description"
        row[2].text = "G1"
        row[3].text = "G2"
        for expr, desc, gate1, gate2 in tr_lst:
            # Adding a row and then adding data in it.
            row = table.add_row().cells
            # Converting id to string as table can only take string input
            row[0].text = expr
            row[1].text = desc
            row[2].text = gate1
            row[3].text = gate2
        table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[2].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[3].paragraphs[0].runs[0].font.bold = True
        for row in table.rows:
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            row.cells[0].paragraphs[0].runs[0].font.name = "Consolas"
        table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
        table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[2].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[3].paragraphs[0].runs[0].font.bold = True
        for cell in table.columns[0].cells:
            cell.width = Cm(7.25)
        for cell in table.columns[1].cells:
            cell.width = Cm(5.75)
        for cell in table.columns[2].cells:
            cell.width = Cm(1.4)
        for cell in table.columns[3].cells:
            cell.width = Cm(1.4)

    doc.save(path + filename_output)


def main():
    system("cls")
    path = "projects/CEOD/Hon/drain_seq/"
    if True:
        filenames = glob(path + "*.cnf.xml")

        for filename_long in filenames:
            filename = filename_long[len(path) :]

            # filename = "DRAIN_6.cnf.xml"
            # filename = "30KBS070_MEPROX.cnf.xml"
            # filename_output_excel = filename.replace(".cnf.xml", ".xlsx")
            filename_output_word = filename.replace(".cnf.xml", ".docx")

            dprint(f"reading sequence {filename}", "GREEN")
            df_TR, df_STEP, TR_dict, STEP_dict = get_sequence(path, filename)
            # dprint(f"writing sequence to excel", "GREEN")
            # write_sequence_excel(df_TR, df_STEP, path, filename_output_excel)
            dprint(f"writing sequence to word", "GREEN")
            write_sequence_word(TR_dict, STEP_dict, path, filename_output_word)
        dprint(f"DONE!", "IMPORTANT")
    else:
        filename = "30UZ530ST.cnf.xml"
        filename_output_excel = filename.replace(".cnf.xml", ".xlsx")
        filename_output_word = filename.replace(".cnf.xml", ".docx")

        dprint(f"reading sequence {filename}", "GREEN")
        df_TR, df_STEP, TR_dict, STEP_dict = get_sequence(path, filename)
        # dprint(f"writing sequence to excel", "GREEN")
        # write_sequence_excel(df_TR, df_STEP, path, filename_output_excel)
        dprint(f"writing sequence to word", "GREEN")
        write_sequence_word(TR_dict, STEP_dict, path, filename_output_word)


if __name__ == "__main__":
    main()
